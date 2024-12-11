import streamlit as st
import pandas as pd
from streamlit_feedback import streamlit_feedback
from src.LLM_source_code.LLM_HumanFeedbackLoop_v2.query import get_expected_response
from src.LLM_source_code.LLM_HumanFeedbackLoop.model import Azure_model_for_human_in_the_loop
from src.LLM_source_code.LLM_HumanFeedbackLoop.annotation_llm import assistant_call_with_annotation
from src.LLM_source_code.LLM_HumanFeedbackLoop_v2.reset import Update_vector
from src.LLM_source_code.LLM_HumanFeedbackLoop_v2.LLM_Langchain_Call import langchain_rag
from src.LLM_source_code.LLM_HumanFeedbackLoop_v2.LLM_Eval_Model import multiturn_generate_content_rel,multiturn_generate_content_indirect,multiturn_generate_correctness,multiturn_generate_faithfull
import os
import time
import re
import shutil
from docx import Document
from src.LLM_source_code.LLM_HumanFeedbackLoop_v2.LLM_Langchain_Call import append_to_vectorstore,reset_vectorstore

def updated_text_based():
    w1, col1, w2, col2, w3 = st.columns([1, 5, 1, 5, 1]) 
    m1, m2, m3 = st.columns([0.2, 8, 0.2])
    m4, m5, m6 = st.columns([0.2, 8, 0.2])
    # Update_vector()
    # Row 1: Select Use Case
    with col1:
        st.markdown("<div style='height:4rem; line-height:4rem; font-weight: bold;'>Select Use Case</div>", unsafe_allow_html=True)
    with col2:
        vAR_usecase = st.selectbox(" ", ("Policy Guru", "Multimodal RAG"))

    # Row 2: Select LLM
    with col1:
        st.markdown("<div style='height:5.3rem; line-height:5.3rem; font-weight: bold;'>Select LLM</div>", unsafe_allow_html=True)
    with col2:
        vAR_model = st.selectbox(" ", ("All","gpt-4o(Azure OpenAI)", "Langchain(OpenAI)"  ))

    # Row 3: Select Platform
    with col1:
        st.markdown("<div style='height:4rem; line-height:4rem; font-weight: bold;'>Select Platform</div>", unsafe_allow_html=True)
    with col2:
        if vAR_model == "All":
            vAR_platform = st.selectbox(" ", ("All", "Assistant(Azure OpenAI)", "AWS Bedrock", 
                                             "Vertex AI(Gemini)", "Azure OpenAI(Langchain)"))
        elif vAR_model in ["gpt-4o(Azure OpenAI)", "gpt-4(Azure OpenAI)"]:
            vAR_platform = st.selectbox(" ", ("Assistant(Azure OpenAI)",))
        elif vAR_model == "claude-3.5-sonnet(Bedrock)":
            vAR_platform = st.selectbox(" ", ("AWS Bedrock",))
        elif vAR_model == "gemini-1.5(Vertex AI)":
            vAR_platform = st.selectbox(" ", ("Vertex AI(Gemini)",))
        elif vAR_model == "Langchain(OpenAI)":
            vAR_platform = st.selectbox(" ", ("Langchain",))

    # Row 4: Select LLM as a Judge Model
    with col1:
        st.markdown("<div style='height:5.1rem; line-height:5.1rem; font-weight: bold;'>Select LLM as a Judge Model (Evaluator)</div>", 
                    unsafe_allow_html=True)
    with col2:
        vAR_eval_llm = st.selectbox(" ", ("Gemini", "GPT", "Claude"))

    # # Row 5: Select LLM as a Judge Type
    # with col1:
    #     st.markdown("<div style='height:4rem; line-height:4rem; font-weight: bold;'>Select LLM as a Judge Type</div>", unsafe_allow_html=True)
    # with col2:
    #     vAR_eval_type = st.selectbox(" ", ("Pairwise Comparison(Default)", 
    #                                       "Evaluation by criteria with reference", 
    #                                       "Evaluation by criteria without reference"))
    with m2:
        # Initialize chat history and feedback state
        if "messages" not in st.session_state:
            st.session_state.messages = [
                {"role": "user", "content": "We are delighted to have you here in the Live Agent Chat room!"},
                {"role": "assistant", "content": "Hello! How can I assist you today?"}
            ]
        if "feedback_given" not in st.session_state:
            st.session_state.feedback_given = set()
        if "file_copied" not in st.session_state:
            st.session_state.file_copied = False
        if "improved_response" not in st.session_state:
            st.session_state["improved_response"] = None
        if "prompt" not in st.session_state:
            st.session_state["prompt"] = None
        if "first_response" not in st.session_state:
            st.session_state["first_response"] = None
        if "langchain_response" not in st.session_state:
            st.session_state["langchain_response"] = None
        if "expected_response" not in st.session_state:
            st.session_state["expected_response"] = None
        if "langchain_reference" not in st.session_state:
            st.session_state["langchain_reference"] = None
        if "feedback_openai" not in st.session_state:
            st.session_state["feedback_openai"] = None
        if "eval_result_df" not in st.session_state:
            st.session_state["eval_result_df"] = None
        if "improved_response_lgn" not in st.session_state:
            st.session_state['improved_response_lgn'] = None

        # Add custom CSS for styling chat
        st.markdown(
            """
            <style>
            .chat-container { display: flex; align-items: center; margin: 10px 0; }
            .chat-container.user { justify-content: flex-end; }
            .chat-container.assistant { justify-content: flex-start; }
            .chat-container img { width: 40px; height: 40px; border-radius: 50%; margin: 0 10px; }
            .chat-bubble { max-width: 70%; padding: 10px; border-radius: 10px; margin: 5px 0; }
            .chat-bubble.user { background-color: #e0e0e0; text-align: right; }
            .chat-bubble.assistant { background-color: #ffffff; text-align: left; }
            </style>
            """,
            unsafe_allow_html=True,
        )

        user_image_url = "https://storage.googleapis.com/macrovector-acl-eu/previews/118720/thumb_118720.webp"
        assistant_image_url = "https://cdn-icons-png.flaticon.com/512/6014/6014401.png"

        # Display chat history
        for i, message in enumerate(st.session_state.messages):
            if message["role"] == "user":
                st.markdown(
                    f'''
                    <div class="chat-container user">
                        <div class="chat-bubble user">{message["content"]}</div>
                        <img src="{user_image_url}" alt="User">
                    </div>
                    ''',
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'''
                    <div class="chat-container assistant">
                        <img src="{assistant_image_url}" alt="Assistant">
                        <div class="chat-bubble assistant">{message["content"]}</div>
                    </div>
                    ''',
                    unsafe_allow_html=True,
                )
                
                if "table_heading" in message:
                    st.write(f"### {message['table_heading']}")
                if "table" in message:
                    st.table(message["table"])
                if "feedback_table_heading" in message:
                    st.write(f"### {message['feedback_table_heading']}")
                if "feedback_table" in message:
                    st.table(message["feedback_table"])
                if "third_table_heading" in message:  
                    st.write(f"### {message['third_table_heading']}")
                if "third_table" in message:  
                    st.table(message["third_table"])
                
                # Collect feedback for assistant responses
                if i > 1 and i not in st.session_state.feedback_given:
                    selected_option = st.radio("Choose a model below to provide feedback:", options=["OpenAI", "Langchain", "Both"], horizontal=True ,key=f"radio_{i}" )
                    if selected_option == "OpenAI":
                        feedback_ = streamlit_feedback(
                            align="flex-start",
                            feedback_type="thumbs",
                            optional_text_label="[Optional] Please provide an explanation",
                            key=f"thumbs_{i}"
                        )
                        if feedback_:
                            # Parse feedback details
                            feedback_score = feedback_.get("score", "neutral")
                            feedback_text = feedback_.get("text", "")
                            st.session_state.feedback_given.add(i)
                            
                            if feedback_score == "👎":
                                # Handle thumbs-down feedback
                                prompt = st.session_state['prompt']

                                response = st.session_state['first_response']
                                improved_response = Azure_model_for_human_in_the_loop(prompt, response, feedback_score, feedback_text)
                                st.session_state["improved_response"] = improved_response
                                    
                                data_2 = {
                                    "Prompt": [prompt],
                                    "Expected Response": [st.session_state["expected_response"]],
                                    "GPT (Standard)": [response],
                                    "GPT (with Langchain)": [st.session_state["langchain_response"]],
                                    "Label" : [feedback_score],
                                    "Feedback" : [feedback_text] 
                                }
                                step_2 = pd.DataFrame(data_2)
                                step_3 = step_2.copy()
                                step_3["Optimized GPT(STD)"] = improved_response
                                
                                user_message = st.session_state['prompt']
                                assistant_message = st.session_state["improved_response"]
                                if not st.session_state.file_copied:
                                    try:
                                        source_file = "DMV_FAQ.docx"
                                        destination_file = "DMV_FAQ_copy.docx"
                                        shutil.copy(source_file, destination_file)
                                        st.session_state.file_copied = True
                                    except Exception as e:
                                        st.error(f"Error during file operations: {e}")

                                file_path = "DMV_FAQ_copy.docx"
                                text = f"\n\n{user_message}\n\n{assistant_message}"

                                # Update the .docx file
                                doc = Document(file_path)
                                doc.add_paragraph(text)
                                doc.save(file_path)
                                # Update vector store
                                vector_store_files = st.session_state.client.beta.vector_stores.files.list(
                                    vector_store_id=os.getenv("VECTOR_STORE_ID")
                                )
                                if vector_store_files.data:
                                    file_ids = vector_store_files.data[0].id
                                    st.session_state.client.beta.vector_stores.files.delete(
                                        vector_store_id=os.getenv("VECTOR_STORE_ID"), file_id=file_ids
                                    )
                                response = st.session_state.client.files.create(file=open(file_path, "rb"), purpose="assistants")
                                file_id = response.id
                                st.session_state.client.beta.vector_stores.files.create(
                                    vector_store_id=os.getenv("VECTOR_STORE_ID"), file_id=file_id
                                )
                                st.session_state.thread = st.session_state.client.beta.threads.create()
                                

                                # Append improved response and tables with headings
                                st.session_state.messages.append(
                                    {
                                        "role": "assistant",
                                        "content": "",
                                        "table_heading": "Human-in-the-Loop (HITL):",
                                        "table": step_2,
                                        "feedback_table_heading": "Optimized LLM response based on human feedback:",
                                        "feedback_table": step_3,
                                        "third_table_heading": "Model Evaluation:", 
                                        "third_table": st.session_state["eval_result_df"]
                                    }
                                )
                                st.write("User Feedback has been successfully saved to the OpenAI Vector Database")
                                
                            if feedback_score == "👍":
                                st.write("User Feedback has been saved successfully")

                    if selected_option == "Langchain":
                        feedback_ = streamlit_feedback(
                            align="flex-start",
                            feedback_type="thumbs",
                            optional_text_label="[Optional] Please provide an explanation",
                            key=f"thumbs_{i}"
                        )
                        if feedback_:
                            # Parse feedback details
                            feedback_score = feedback_.get("score", "neutral")
                            feedback_text = feedback_.get("text", "")
                            st.session_state.feedback_given.add(i)
                            
                            if feedback_score == "👎":
                                # Handle thumbs-down feedback
                                prompt = st.session_state['prompt']

                                response = st.session_state['first_response']
                                improved_response = Azure_model_for_human_in_the_loop(prompt, response, feedback_score, feedback_text)
                                st.session_state["improved_response"] = improved_response
                                                                    
                                data_2 = {
                                    "Prompt": [prompt],
                                    "Expected Response": [st.session_state["expected_response"]],
                                    "GPT (Standard)": [response],
                                    "GPT (with Langchain)": [st.session_state["langchain_response"]],
                                    "Label" : [feedback_score],
                                    "Feedback" : [feedback_text] 
                                }
                                step_2 = pd.DataFrame(data_2)
                                step_3 = step_2.copy()
                                step_3["Optimized GPT(Langchain)"] = improved_response
                                
                                user_message = st.session_state['prompt']
                                assistant_message = st.session_state["improved_response"] 
                                text = f"\n\n{user_message}\n\n{assistant_message}"
                                append_to_vectorstore(text)

                                # Append improved response and tables with headings
                                st.session_state.messages.append(
                                    {
                                        "role": "assistant",
                                        "content": "",
                                        "table_heading": "Human-in-the-Loop (HITL):",
                                        "table": step_2,
                                        "feedback_table_heading": "Optimized LLM responses based on human feedback:",
                                        "feedback_table": step_3,
                                        "third_table_heading": "Model Evaluation:",  
                                        "third_table": st.session_state["eval_result_df"],
                                    }
                                )
                                st.write("User Feedback has been successfully saved to the Chroma Vector Database")
                            
                            if feedback_score == "👍":
                                st.write("User Feedback has been saved successfully")

                    if selected_option == "Both":
                        feedback_ = streamlit_feedback(
                            align="flex-start",
                            feedback_type="thumbs",
                            optional_text_label="[Optional] Please provide an explanation",
                            key=f"thumbs_{i}"
                        )
                        st.session_state["feedback_openai"] = feedback_
                        
                        if feedback_:
                            # Parse feedback details
                            feedback_score = feedback_.get("score", "neutral")
                            feedback_text = feedback_.get("text", "")
                            st.session_state.feedback_given.add(i)
                            
                            if feedback_score == "👎":
                                # Handle thumbs-down feedback
                                prompt = st.session_state['prompt']

                                response_std = st.session_state['first_response']
                                response_lgn = st.session_state["langchain_response"]
                                improved_response_std = Azure_model_for_human_in_the_loop(prompt, response_std, feedback_score, feedback_text)
                                improved_response_lgn = Azure_model_for_human_in_the_loop(prompt, response_lgn, feedback_score, feedback_text)
                                st.session_state["improved_response"] = improved_response_std
                                st.session_state['improved_response_lgn'] = improved_response_lgn
                                data_2 = {
                                    "Prompt": [prompt],
                                    "Expected Response": [st.session_state["expected_response"]],
                                    "GPT (Standard)": [response_std],
                                    "GPT (with Langchain)": [st.session_state["langchain_response"]],
                                    "Label" : [feedback_score],
                                    "Feedback" : [feedback_text] 
                                }
                                step_2 = pd.DataFrame(data_2)
                                step_3 = step_2.copy()
                                step_3["Optimized GPT(STD)"] = improved_response_std
                                step_3["Optimized GPT(Langchain)"] = st.session_state['improved_response_lgn']
                                
                                user_message = st.session_state['prompt']
                                assistant_message = st.session_state["improved_response"] 
                                if not st.session_state.file_copied:
                                    try:
                                        source_file = "DMV_FAQ.docx"
                                        destination_file = "DMV_FAQ_copy.docx"
                                        shutil.copy(source_file, destination_file)
                                        st.session_state.file_copied = True
                                    except Exception as e:
                                        st.error(f"Error during file operations: {e}")

                                file_path = "DMV_FAQ_copy.docx"
                                text = f"\n\n{user_message}\n\n{assistant_message}"

                                # Update the .docx file
                                doc = Document(file_path)
                                doc.add_paragraph(text)
                                doc.save(file_path)
                                # Update vector store
                                vector_store_files = st.session_state.client.beta.vector_stores.files.list(
                                    vector_store_id=os.getenv("VECTOR_STORE_ID")
                                )
                                if vector_store_files.data:
                                    file_ids = vector_store_files.data[0].id
                                    st.session_state.client.beta.vector_stores.files.delete(
                                        vector_store_id=os.getenv("VECTOR_STORE_ID"), file_id=file_ids
                                    )
                                response = st.session_state.client.files.create(file=open(file_path, "rb"), purpose="assistants")
                                file_id = response.id
                                st.session_state.client.beta.vector_stores.files.create(
                                    vector_store_id=os.getenv("VECTOR_STORE_ID"), file_id=file_id
                                )
                                st.session_state.thread = st.session_state.client.beta.threads.create()
                                # st.write("User Feedback has been successfully saved to the OpenAI Vector Database")
                                append_to_vectorstore(text)
                                
                                # st.session_state["improved_response"] = None

                                # Append improved response and tables with headings
                                st.session_state.messages.append(
                                    {
                                        "role": "assistant",
                                        "content": "",
                                        "table_heading": "Human-in-the-Loop (HITL):",
                                        "table": step_2,
                                        "feedback_table_heading": "Optimized LLM response based on human feedback:",
                                        "feedback_table": step_3,
                                        "third_table_heading": "Model Evaluation:",  # Add heading for the third table
                                        "third_table": st.session_state["eval_result_df"],
                                    }
                                )
                                st.write("User Feedback has been successfully saved to the OpenAI and Chroma Vector Database")
                            if feedback_score == "👍":
                                st.write("User Feedback has been saved successfully")

    # Capture user input
    with m2:
        prompt = st.chat_input("What else can I do to help?")
    if prompt:
        st.session_state["prompt"] = prompt
        # Append user message
        st.session_state.messages.append({"role": "user", "content": prompt})

        # Generate assistant response with a table
        response_text = assistant_call_with_annotation(prompt)
        st.session_state["first_response"] = response_text
        project_id = "elp-2022-352222"
        dataset_id = "DMV_ELP"
        table_name = "prompt_and_response"

        expected_response = get_expected_response(project_id, dataset_id, table_name, prompt)
        st.session_state["expected_response"] = expected_response
        langchain_response, langchain_reference = langchain_rag(st.session_state["prompt"])
        st.session_state["langchain_reference"] = langchain_reference

        st.session_state["langchain_response"] = langchain_response
        data = {
            "Prompt": [st.session_state["prompt"]],
            "Expected Response": [st.session_state["expected_response"]],
            "GPT (Standard)": [st.session_state["first_response"]],
            "GPT (with Langchain)": [st.session_state["langchain_response"]]
        }
        step_1 = pd.DataFrame(data)
        data_eval_lang = {
            "input_prompt": [st.session_state["prompt"]],
            "response": [st.session_state["first_response"]],
            "reference_text": [st.session_state["langchain_reference"]],
            "model":["GPT (with Langchain)"]
        }
        data_eval_1 = pd.DataFrame(data_eval_lang)
        data_eval_openai = {
            "input_prompt": [st.session_state["prompt"]],
            "response": [st.session_state["first_response"]],
            "reference_text": [st.session_state["langchain_reference"]],
            "model":["GPT (Standard)"]
        }
        data_eval_2 = pd.DataFrame(data_eval_openai)
        
        vAR_eval_df1_rel = multiturn_generate_content_rel(data_eval_1)
        vAR_eval_df1_indirect = multiturn_generate_content_indirect(data_eval_1)
        vAR_eval_df1_faith = multiturn_generate_faithfull(data_eval_1)
        vAR_eval_df1_correctness = multiturn_generate_correctness(data_eval_1)

        vAR_eval_df2_rel = multiturn_generate_content_rel(data_eval_2)
        vAR_eval_df2_indirect = multiturn_generate_content_indirect(data_eval_2)
        vAR_eval_df2_faith = multiturn_generate_faithfull(data_eval_2)
        vAR_eval_df2_correctness = multiturn_generate_correctness(data_eval_2)

        concatenated_df = pd.concat([vAR_eval_df1_rel, vAR_eval_df1_indirect,vAR_eval_df1_faith,vAR_eval_df1_correctness,vAR_eval_df2_rel,vAR_eval_df2_indirect,vAR_eval_df2_faith,vAR_eval_df2_correctness], ignore_index=True)
        st.session_state["eval_result_df"] = concatenated_df
        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": "",
                "table_heading": "Model Response:",
                "table": step_1
            }
        )
        st.rerun()
    
    
    # if st.session_state["langchain_reference"] != None:
    #     if st.session_state["langchain_reference"] != st.session_state["present_langchain_reference"] or st.session_state["improved_response"] != None:
    #         data_eval_lang = {
    #             "input_prompt": [st.session_state["prompt"]],
    #             "response": [st.session_state["first_response"]],
    #             "reference_text": [st.session_state["langchain_reference"]],
    #             "model":["GPT (with Langchain)"]
    #         }
    #         data_eval_1 = pd.DataFrame(data_eval_lang)
    #         data_eval_openai = {
    #             "input_prompt": [st.session_state["prompt"]],
    #             "response": [st.session_state["first_response"]],
    #             "reference_text": [st.session_state["langchain_reference"]],
    #             "model":["GPT (Standard)"]
    #         }
    #         data_eval_2 = pd.DataFrame(data_eval_openai)
            
    #         vAR_eval_df1_rel = multiturn_generate_content_rel(data_eval_1)
    #         vAR_eval_df1_indirect = multiturn_generate_content_indirect(data_eval_1)
    #         vAR_eval_df2_rel = multiturn_generate_content_rel(data_eval_2)
    #         vAR_eval_df2_indirect = multiturn_generate_content_indirect(data_eval_2)
            
    #         concatenated_df = pd.concat([vAR_eval_df1_rel, vAR_eval_df1_indirect,vAR_eval_df2_rel,vAR_eval_df2_indirect], ignore_index=True)
    #         st.session_state.messages.append({"role": "assistant", "content": "", "table_heading": "Model Evaluation:","table": concatenated_df})
    #         st.session_state["present_langchain_reference"] = st.session_state["langchain_reference"]
    #         st.rerun()
        
    with m2:
        if st.button("Reset VectorDB"):
            try:
                # st.session_state.clear()
                Update_vector()
                reset_vectorstore("DMV_FAQ.docx")
                st.session_state.thread = st.session_state.client.beta.threads.create()
                st.write("KnowledgeBase successfully reset")
                time.sleep(2)
                st.rerun()
            except PermissionError:
                st.warning("There is a problem resetting the VectorDB. Kindly refresh and try again.")
            
            

